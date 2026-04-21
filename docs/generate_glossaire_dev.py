#!/usr/bin/env python3
"""Génère le glossaire développement au format .docx à partir de glossaire_data.json."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, json
from datetime import datetime

HERE = os.path.dirname(os.path.abspath(__file__))
DATA = os.path.join(HERE, "glossaire_data.json")
OUTPUT = "/Users/mathieuvirondaud/claude/Glossaire_Developpement.docx"

DARK_GREEN = RGBColor(0x00, 0x60, 0x39)
GREEN = RGBColor(0x16, 0xa3, 0x4a)
GREY = RGBColor(0x64, 0x74, 0x8b)
BLUE = RGBColor(0x08, 0x91, 0xb2)


def h1(doc, text, page_break=True):
    if page_break:
        doc.add_page_break()
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'E6F4EA')
    cell._tc.get_or_add_tcPr().append(shd)
    r = cell.paragraphs[0].add_run(text)
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = DARK_GREEN
    doc.add_paragraph()


def para(doc, text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color


def term_box(doc, emoji, term, english, definition, example, color='E6F4EA'):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shd)
    p1 = cell.paragraphs[0]
    r = p1.add_run(f"{emoji}  {term}")
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = DARK_GREEN
    if english:
        r2 = p1.add_run(f"   ({english})")
        r2.italic = True; r2.font.size = Pt(11); r2.font.color.rgb = GREY
    p2 = cell.add_paragraph()
    r = p2.add_run("En français : "); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = BLUE
    r2 = p2.add_run(definition); r2.font.size = Pt(11)
    p3 = cell.add_paragraph()
    r = p3.add_run("Dans ton projet : "); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = GREEN
    r2 = p3.add_run(example); r2.italic = True; r2.font.size = Pt(10)
    doc.add_paragraph()


def main():
    with open(DATA, encoding='utf-8') as f:
        data = json.load(f)

    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2); s.bottom_margin = Cm(2)
        s.left_margin = Cm(2); s.right_margin = Cm(2)

    # Couverture
    meta = data['meta']
    hero = meta.get('hero_image')
    if hero and os.path.exists(hero):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(hero, width=Inches(6.0))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("📖"); r.font.size = Pt(50)

    t = doc.add_heading(meta['title'], level=0)
    for r in t.runs: r.font.color.rgb = DARK_GREEN; r.font.size = Pt(28)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(meta['subtitle']); r.italic = True; r.font.size = Pt(12); r.font.color.rgb = GREY

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("━━━━━━━━━━━━━━━━━━━━"); r.font.size = Pt(8); r.font.color.rgb = GREEN

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Chaque terme est expliqué simplement,\navec un exemple tiré de tes projets.")
    r.font.size = Pt(11); r.font.color.rgb = GREY

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{datetime.now().strftime('%d/%m/%Y')}  •  {meta['author']}")
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = DARK_GREEN

    # Sections de termes (1-6 et 10)
    for section in data['sections']:
        h1(doc, f"{section['num']}. {section['title']}")
        if section.get('intro'):
            para(doc, section['intro'], italic=True)
        for term in section['terms']:
            term_box(doc, term['emoji'], term['term'], term.get('english', ''),
                     term['definition'], term['example'], color=section['color'])

    # Section 7 — table projets
    pt = data['projects_table']
    h1(doc, f"{pt['num']}. {pt['title']}")
    para(doc, pt['intro'], bold=True)
    para(doc, "")
    tbl = doc.add_table(rows=len(pt['rows']) + 1, cols=3)
    tbl.style = 'Light Grid Accent 1'
    hdr = tbl.rows[0].cells
    hdr[0].text = "Projet"; hdr[1].text = "Rôle"; hdr[2].text = "Statut"
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True; r.font.color.rgb = DARK_GREEN
    for i, row in enumerate(pt['rows'], 1):
        for j, val in enumerate(row):
            tbl.rows[i].cells[j].text = val
    doc.add_paragraph()
    para(doc, pt['note'], italic=True)

    # Section 8 — commandes
    cm = data['commands']
    h1(doc, f"{cm['num']}. {cm['title']}")
    para(doc, cm['intro'], italic=True)
    for title, cmd, desc in cm['items']:
        para(doc, title, bold=True, size=12, color=DARK_GREEN)
        p = doc.add_paragraph()
        r = p.add_run("  " + cmd)
        r.font.name = 'Courier New'; r.font.size = Pt(10); r.font.color.rgb = BLUE
        para(doc, "  → " + desc, italic=True, size=10)
        para(doc, "")

    # Section 9 — à retenir
    rc = data['recap']
    h1(doc, f"{rc['num']}. {rc['title']}")
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'D1FAE5')
    cell._tc.get_or_add_tcPr().append(shd)
    for icon_term, role in rc['items']:
        p = cell.add_paragraph()
        r = p.add_run(f"▸ {icon_term}  ")
        r.bold = True; r.font.size = Pt(13); r.font.color.rgb = DARK_GREEN
        r2 = p.add_run(f"= {role}"); r2.font.size = Pt(11)

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Pas besoin de tout retenir.\nCe document reste ici, tu le consultes quand tu en as besoin.")
    r.italic = True; r.font.size = Pt(12); r.font.color.rgb = GREY

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("⛳ Guide du développeur occasionnel — Mathieu Virondaud")
    r.bold = True; r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GREEN

    doc.save(OUTPUT)
    print(f"✅ Document créé : {OUTPUT}")
    print(f"📄 Taille : {os.path.getsize(OUTPUT):,} octets")


if __name__ == "__main__":
    main()
